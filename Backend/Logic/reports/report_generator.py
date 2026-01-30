import json
import os

from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert
from pathlib import Path

# TODO: replace with MongoDB code
FILEPATH = Path("../analysis_ready")
SCORE = Path("../scoring/friendliness_summary.json")
OUTPUT = Path("generated_reports")

seen = set()


def generate_report(file_path, score_path, output_path, pdf=False):
    # Open the extracted text and scoring files.
    file_path = Path(file_path)
    score_path = Path(score_path)
    output_path = Path(output_path)
    with open(file_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    with open(score_path, "r", encoding="utf-8") as f:
        scores = json.load(f)

    # Get the file name to search for in the score summary.
    filename = Path(data.get("file", "unnamed file")).with_suffix(".json").name

    score = scores.get(filename, {}).get("foodtruck", "N/A")

    keyword_contexts = data.get("keyword_contexts", {})

    # Create the file and add the heading and score.
    doc = Document()
    doc.add_heading(f"Summary Report for {data.get('file', 'unnamed file')}", level=1)

    doc.add_paragraph(f"Overall Score: {score}%")

    # Formatting for the table.
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Category"
    hdr_cells[1].text = "Status"

    # Put a Green for a found category and Red for a missing category.
    for category, content in keyword_contexts.items():
        empty = (len(content) == 0)
        status = "Found" if not empty else "Missing"
        row_cells = table.add_row().cells
        row_cells[0].text = category.capitalize()
        row_cells[1].text = status
        tc = row_cells[1]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        if empty:
            shd.set(qn('w:fill'), 'A61B00')
        else:
            shd.set(qn('w:fill'), '078701')
        tcPr.append(shd)

    for row in table.rows:
        row.cells[0].width = Inches(3)
        row.cells[1].width = Inches(1)

    doc.add_heading('Key Findings', level=1)
    seen = set()
    for category, content in keyword_contexts.items():
        for subcategory, items in content.items():
            if isinstance(items, list):
                for line in items:
                    if line not in seen and len(line.split()) >= 4:
                        if "." not in line:
                            doc.add_paragraph((line + ".").capitalize())
                        elif "$" in line:
                            doc.add_paragraph((line + ".").capitalize())
                        else:
                            doc.add_paragraph(line.capitalize())
                        seen.add(line)

    doc.add_heading('Recommendations', level=1)
    for category, content in keyword_contexts.items():
        empty = (len(content) == 0)
        if empty:
            doc.add_paragraph(f"Find more information for {category}.")

    doc.save(output_path)

    pdf_path = None

    if pdf:
        pdf_path = output_path.with_suffix(".pdf")
        convert(output_path, pdf_path)
        print(f"PDF saved to: {pdf_path}")

    return output_path, pdf_path


if __name__ == "__main__":
    OUTPUT.mkdir(exist_ok=True)
    for file_path in FILEPATH.iterdir():
        output_path = OUTPUT / f"{file_path.stem}_Report.docx"
        generate_report(file_path, SCORE, output_path, pdf=False)
