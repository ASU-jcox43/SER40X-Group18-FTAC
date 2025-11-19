import json
import os

from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert

FILEPATH = "../analysis_ready"
SCORE = "../scoring/friendliness_summary.json"
OUTPUT = "generated reports"

seen = set()


def json_to_table(file_path, score_path, output_path, pdf=False):
    # Open the extracted text and scoring files.
    with open(file_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    with open(score_path, "r", encoding="utf-8") as f:
        scores = json.load(f)

    # Get the file name to search for in the score summary.
    filename = data.get("file", "unnamed file")
    filename = filename.replace(".txt", ".json")
    score = scores.get(filename, "N/A")

    keyword_contexts = data.get("keyword_contexts", {})

    # Create the file and add the heading and score.
    doc = Document()
    doc.add_heading(f"Summary Report for {data.get('file', 'unnamed file')}", level=1)

    doc.add_paragraph(f"Overall Score: {score}%")

    # Formatting for the table.
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Category"
    hdr_cells[1].text = "Found"

    # Put a Green for a found category and Red for a missing category.
    for category, content in keyword_contexts.items():
        empty = (len(content) == 0)
        status = "+" if not empty else "-"
        row_cells = table.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = status
        tc = row_cells[1]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        if empty:
            shd.set(qn('w:fill'), 'A61B00')
        else:
            shd.set(qn('w:fill'), '078701')
        tcPr.append(shd)

    for row in table.rows:
        row.cells[0].width = Inches(3)
        row.cells[1].width = Inches(1)

    doc.add_heading('Key Findings', level=1)
    for category, content in keyword_contexts.items():
        for subcategory, items in content.items():
            if isinstance(items, list):
                for line in items:
                    if line not in seen:
                        doc.add_paragraph(line)
                        seen.add(line)

    doc.add_heading('Recommendations', level=1)
    for category, content in keyword_contexts.items():
        empty = (len(content) == 0)
        if empty:
            doc.add_paragraph(f"Find more information for {category}.")

    doc.save(output_path)
    print(f"Word document saved to: {output_path}")

    if pdf:
        pdf_path = output_path.replace(".docx", ".pdf")
        convert(output_path, pdf_path)
        print(f"PDF saved to: {pdf_path}")


if __name__ == "__main__":
    for filename in os.listdir(FILEPATH):
        file_path = os.path.join(FILEPATH, filename)
        output_name = os.path.splitext(filename)[0] + "_Report.docx"
        output_path = os.path.join(OUTPUT, output_name)
        json_to_table(file_path, SCORE, output_path, pdf=True)
